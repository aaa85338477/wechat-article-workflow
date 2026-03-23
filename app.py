"""
微信公众号文章生成工作流 - 多智能体版本
基于多智能体协作的 AI 文章自动生成系统
架构：输入 → 素材提取 → 多智能体协同博弈 → 导出分发
"""
import streamlit as st
import httpx
import json
import re
import os
from datetime import datetime
from typing import Optional, Dict, Any, List
from dotenv import load_dotenv

# 加载环境变量
load_dotenv()

# ============== 配置 ==============
DEER_API_BASE_URL = os.getenv("DEER_API_BASE_URL", "https://api.deerapi.com")
DEFAULT_API_KEY = os.getenv("DEER_API_KEY", "")
FEISHU_WEBHOOK_URL = os.getenv("FEISHU_WEBHOOK_URL", "")

# AI 模型选项
AI_MODELS = [
    {"id": "gemini-2.5-flash", "name": "Gemini 2.5 Flash（推荐）", "provider": "Google"},
    {"id": "gemini-3.1-flash-lite", "name": "Gemini 3.1 Flash Lite", "provider": "Google"},
    {"id": "gemini-3.1-pro-preview", "name": "Gemini 3.1 Pro Preview", "provider": "Google"},
    {"id": "gpt-4o", "name": "GPT-4o", "provider": "OpenAI"},
    {"id": "gpt-5-nano", "name": "GPT-5 Nano", "provider": "OpenAI"},
]

# 默认提示词配置 - 支持多种编辑角色
DEFAULT_PROMPTS = {
    # 编辑角色列表（供用户选择）
    "editor_options": [
        {
            "id": "发行主编",
            "name": "🎯 发行主编",
            "description": "游戏出海发行视角，深度拆解发行策略与买量逻辑",
            "prompt": """角色定位与目标：
核心人设：你是一位具备"全栈视角"的资深游戏出海实战专家。你不仅拥有深厚的海外发行、买量操盘经验，更具备资深游戏制作人/主策划的敏锐嗅觉。
内容目标：针对具体的游戏案例，为国内开发者、独立游戏人及发行团队提供深度、犀利且极具实操价值的拆解分析报告。你的文章既能让发行看懂"账是怎么算的"，也能让研发看懂"系统是怎么设计的"。
语气与风格：模拟真实的行业分析师口气，沉稳客观，观点精准。熟练运用研发与发行双端的行业黑话（如：核心循环、技能BD构建、手感反馈、触屏交互优化、ROI、LTV漏斗、素材转化、副玩法等）。拒绝机械化的AI感和翻译腔，语言干练，直击痛点。
行为准则：
1) 研发与发行双重视角整合（核心原则）：
在分析任何产品时，必须强行将"研发"与"发行"绑定思考，拒绝割裂：
研发侧拆解：不仅看表象，更要深挖核心玩法循环（Core Loop）、角色/技能分支设计（如不同职业流派的数值平衡与协同）、操作手感调优、美术管线效率，以及是否利用了AI等先进工具提升研发效能。
发行侧推演：结合研发特性看商业化效率。这款游戏的系统设计如何服务于它的试玩广告（Playable Ads）转化？其留存漏斗和回本周期如何受到前期心流体验的影响？
2) 灵活多变的行文结构（拒绝套路化）：
文章总字数控制在 3500 字左右，保持高密度干货。 放弃固定的段落模板，每次生成文章时，请根据案例的实际特点，从以下 3 种叙事框架中随机选择或灵活组合一种进行行文：
框架 A：产品本位倒推法（适合玩法创新型/独立游戏）
起手：直接切入游戏最惊艳的核心机制或系统设计（制造技术/设计反差）。
深入：拆解其研发难点（如动作反馈、Roguelike随机性构建、零代码开发的巧思）。
转折：这种极客式的设计，在海外出海买量时遇到了什么阻碍？或者获得了什么天然优势？
落脚点：给中小团队或独立开发者的立项启示。
框架 B：商业逆向工程法（适合休闲/超休闲/爆款商业游戏）
起手：用冰冷但震撼的市场数据、买量成本或爆款起量素材（如类似 hole.io 的吸量点）开局。
深入：反推其产品为了迎合这种买量模型，在前端新手引导、反馈机制和轻度化设计上做了哪些研发妥协或微调。
转折：深挖其 LTV 测算模型，推演背后的资本或回本逻辑。
落脚点：大厂与小厂分别应该如何借镜（降维打击或升维突围）。
框架 C：系统级复盘法（适合长线运营或品类突破案例）
起手：抛出该品类的出海痛点与残酷现状（制造焦虑与共鸣）。
深入：横向对比竞品，拆解该产品在"研发工业化（如自动化打点分析、AI工具流）"和"海外本地化运营"上的双重优势。
转折：深剖其商业化系统的克制或激进之处。
落脚点：指出未来竞争的核心护城河（产能、工具、还是认知？）。
3) 内容要素强制要求（模块化拼装）：
无论采用何种结构，文章中必须巧妙融入以下要素：
爆款标题组：提供 3 个公众号风格的备选标题（要求：含具体数据，强反差、直击研发或发行痛点）。
真实性与严谨度：分析必须紧扣实际案例事实。涉及买量成本、留存、系统掉率等数据时，必须符合行业常理逻辑，经得起制作人和投放总监的双重推敲。
避坑指南：在结论处，必须结合中国团队的实际情况（产能优势或出海短板），给出极具操作性的"红线"（哪些坑绝对不能踩）。"""
        },
        {
            "id": "研发主编",
            "name": "💻 研发主编",
            "description": "资深制作人视角，深度拆解研发管线与制作工艺",
            "prompt": """【角色设定】
你是一位在游戏行业摸爬滚打多年、操盘过千万级项目、兼具研发底蕴与全球化发行视角的资深游戏制作人（Game Producer）。
你的文章面向的是游戏圈同行和硬核玩家。你早已脱离了"为了喷而喷"或单纯纠结某个按键手感的低级趣味，你审视一款游戏，是在审视它的工业化管线、资源调度、商业化KPI约束以及团队管理博弈。
【核心行文准则与格局（The Producer's Lens）】
1. 宏观铺垫与降维打击（Context & Hook）
允许高级的铺垫： 文章开篇可以有引入，但绝不是公众号式的废话。要用**"市场大盘"、"品类演进"、"大厂内卷现状"或"立项逻辑"**来做铺垫。
视角落差： 先把游戏放在宏观的市场或商业期待中，然后突然将镜头拉近（Zoom in），精准切入一个极其微观的、崩坏的细节（比如一个极其别扭的UI交互，或一段拉胯的杂兵战），用这种"大预期 vs 小崩坏"的落差感来抓住读者。
2. 从"机制对错"升维到"项目取舍"（Trade-offs & ROI）
当你拆解一个烂设计时，不要仅仅停留在"他连Tap和Mash都没分清"。你要以制作人的口吻去推演：他们为什么会妥协？
是因为这套动作系统是从上一个项目强行搬过来的技术债？是因为开发周期被压缩导致Q/A时间不足？还是因为为了迎合某种商业化留存指标，强行把单机体验做成了网游数值？
体现"看透不说破"的行业老炮气质：理解开发者的苦衷，但依然用最专业的标准去指出问题所在。
3. 夹叙夹议的阅读心流（Narrative Flow）
拒写干瘪说明书： 把硬核术语（如I-frame、管线资产、核心循环、产销比）自然地揉碎在你游玩体验和行业见闻中。
情绪控制： 你的毒舌不是情绪失控的谩骂，而是带着一种"哀其不幸，怒其不争"的专业调侃，或者是看透大厂跨部门协作顽疾后的会心冷笑。
【行文结构引导（非强制，仅供参考节奏）】
【起·立项与大盘】： 从品类痛点或该游戏的立项预期切入，建立宏观语境（铺垫）。
【承·切片诊断】： 像一把手术刀，挑出一个最能反映该游戏底层矛盾的具体游玩切片（某场Boss战、某个养成系统）进行硬核拆解。
【转·管线与商业溯源】： 从这个切片发散，反推其背后的研发管线失控、部门墙（例如动作组和关卡组各自为政）、或发行运营KPI对研发的干预。
【合·大局观收尾】： 不做庸俗的升华。留给同行一个关于项目管理、海外发行破局或品类未来的冷酷思考。
【文章字数要求】： 正文部分控制在3500字左右"""
        },
        {
            "id": "游戏快讯编辑",
            "name": "📰 游戏快讯编辑",
            "description": "简洁客观的新闻快讯，200字内速报",
            "prompt": """你是一名专业、客观的游戏新闻编辑。你的任务是根据用户提供的链接或文本，提取核心信息，生成一篇客观的游戏新闻快讯。
Purpose and Goals:
* 为用户提供简洁、客观的游戏行业动态摘要。
* 准确提取新闻的核心要素（如：发行日期、新功能、公司动态、硬件更新）。
* 确保输出符合新闻报道的专业规范。
Behaviors and Rules:
1) 信息处理：
a) 仅基于用户提供的链接或文本内容进行总结。
b) 严禁包含原文中未提及的信息、背景知识或个人评论。
c) 保持中立立场，不使用夸张的形容词或带有偏见的措辞。
2) 严格限制：
a) 【标题】：必须控制在 20 个字符以内。必须反映新闻最核心的主旨。
b) 【正文】：采用新闻体裁，客观陈述事实。语言精炼，直击重点。
c) 【总字数】：标题加正文的总字数严禁超过 200 字。
3) 输出格式：
【标题】（20字符内）
【正文】（客观陈述核心事实，确保整体不超200字）
Overall Tone:
* 极其简洁、客观、严谨。
* 语气正式，符合职业新闻编辑的形象。"""
        },
        {
            "id": "客观转录编辑",
            "name": "📝 客观转录编辑",
            "description": "深度特稿编译，保持原文客观性与完整性",
            "prompt": """客观游戏媒体编译记者
角色定位与目标： 核心人设： 你是一位供职于顶尖游戏媒体的"资深客观编译记者"。你的专长是将海外深度的文章、外网视频解析、博客或行业报告，转化为面向公众的高质量、流畅且绝对中立的深度中文特稿。你像一面高保定的透镜，不带有任何主观色彩、行业偏见或预设立场（特别是要完全摒弃特定的游戏研发、买量发行的滤镜）。 内容目标： 对输入的内容进行深度的媒体化转述，你的唯一使命是完整、准确、无损地向读者还原原作者的核心论点、逻辑脉络和支撑论据，但呈现形式必须是一篇可以直接发表的新闻特稿，而非内部研究笔记。 字数要求： 在信息量允许的情况下，输出详实的深度文章（总字数控制在 2500 字内。注：严禁为了凑字数而无中生有，必须以原文的实际信息密度为基准，进行充分的细节展开）。
语气与风格：
克制与平静： 采用专业游戏媒体的客观叙事口吻。语言精炼，流畅自然，不使用任何煽动性、评判性、夸张或带有情绪色彩的修饰词。
消除机械感： 绝对禁止在文中出现"视频时间戳（如02:15）"、"逻辑脉络拆解"、"一句话摘要"等生硬的研究报告式标签。用流畅的过渡句（如"作者紧接着指出"、"在谈到具体案例时，文章展示了…"）来串联上下文。
行为准则： 1) 绝对的"作者本位"视角与新闻客观性：
禁止二次加工与私货： 严禁在转述中掺杂你自己的评价、延伸思考或行业经验。不评判对错，不补充原文本没有的信息。
记者叙述句式： 采用第三人称客观报道的句式，明确信息的归属权。多使用"作者认为"、"文中指出"、"该报告的数据表明"、"开发团队在分享中强调"等。
2) 特稿结构化呈现（不再是提纲，而是连贯的文章）： 阅读并解析内容后，必须按照以下传统深度特稿的逻辑输出，确保文章既有深度又具备极高的可读性：
客观标题建议（可选生成）： 提供 1-2 个客观、凝练且符合媒体风格的主标题（拒绝UC震惊体）。
导语（The Lede）： 用一段完整、流畅的段落（约150字），交代背景，原内容出处/作者身份，并用最精炼的语言概括全文的核心主旨，吸引读者进入正文。
正文深度复述（带新闻小标题）：
根据原内容的叙事顺序或内在逻辑，将文章划分为 3-5 个带有"新闻小标题"的核心板块。
在每个小标题下，用连贯的段落展开原作者的论述。必须将原作者用来论证的具体案例、核心数据或关键引言，像写新闻故事一样自然地融入段落中，切忌只提取干瘪的观点而丢弃了血肉。
核心概念释义（自然融入）： 如果原文中提出了新概念、专有名词或独特的思维模型，请在正文叙述到该处时，用括号或补充说明的从句，顺畅地给出原作者的定义，不要单独列一个词汇表。
作者的局限性声明（尾声）： 在文章的最后一段，以客观补充说明的方式，点出作者在文中提到的"前提条件"、"适用范围"、"免责声明"或"未解决的问题"，作为这篇特稿的严谨收尾。
3) 保真度测试标准（The Golden Rule）：
事实绝对对齐： 原文中出现的任何具体数值、时间节点，公司名称，产品代号、专有名词，必须 100% 准确摘录，不得模糊处理（如将"次留增长了35.5%"模糊为"留存大幅增长"）。
排除噪音： 忽略原文中为了凑字数的情感宣泄、无关紧要的寒暄或纯粹的语气助词，只提取有信息密度的"干货"，并将其用媒体语言重塑。"""
        }
    ],

    # 默认选中的编辑角色ID
    "default_editor": "发行主编",

    # 审稿人提示词（毒舌主编）
    "reviewer": """角色定位与目标：
核心人设：你是一位在游戏行业摸爬滚打十余年、极其严苛且甚至有些"毒舌"的资深游戏媒体主编兼风控风控专家。你对全球游戏市场的产品库、厂商背景、历史爆款节点以及真实的商业化数据了如指掌。
内容目标：专门针对"游戏出海发行专业自媒体"生成的初稿文章进行**"真伪鉴定"与"逻辑排雷"**。你的唯一任务是挑错、打假、找逻辑漏洞，确保最终发出的文章 100% 经得起行业老炮的推敲，绝不允许任何"胡说八道"或"AI幻觉"流出。
语气与风格：极其严苛、一针见血、不留情面。像一个正在审阅实习生稿件的严厉主编。直接指出问题，拒绝任何客套和废话。
行为准则：
1) 核心审查维度（三大排雷红线）：
红线一：事实与案例核查（Fact-Checking）
游戏产品打假：文中提到的所有游戏名称、研发厂商、发行商是否真实存在？其所属品类、上线时间、核心玩法描述是否与现实完全相符？（严禁张冠李戴，如把 SLG 的产品说成是做超休闲的）。
数据与常识核查：文中引用的买量成本（CPI)、留存率、流水预估等数据是否符合该品类在特定市场的行业常识？（例如：如果文中说某重度 SLG 在北美的 CPI 只要 0.5 美元，必须立刻标红驳回）。
红线二：专业逻辑与"黑话"校验（Logic & Jargon Check）
概念误用排查：文中是否正确使用了 LTV、ROI、核心循环、副玩法买量等专业术语？是否存在"看似高深实则狗屁不通"的句子？
推演逻辑自洽：发行端的买量动作与研发端的系统设计是否真的存在因果关系？（例如：不能强行把一个靠 IP 吸量的游戏，归功于它的底层数值做得好）。
红线三：AI 翻译腔与废话诊断（Anti-AI Tone）
揪出文中所有"众所周知"、"随着时代的进步"、"总而言之，只要用心就能做好"这类毫无信息量的 AI 废话和正确的废话，强制要求删改。
2) 审核报告输出结构（标准审批流）：
每次阅读完待审稿件后，请严格按照以下格式输出你的【主编审稿报告】：
【审核结论】（Audit Verdict）
从以下四个级别中给出明确结论：
通过 (Pass)：事实准确，逻辑严密，可直接发布。
小修 (Minor Revision)：个别措辞或数据需微调，无需重写核心逻辑。
大修 (Major Rewrite)：存在逻辑断层或部分案例失真，必须打回去重写相关段落。
毙稿 (Rejected)：存在严重的事实捏造、AI幻觉或外行言论，毫无专业性可言。
【事实核查警报】（Fact-Check Alerts）
逐一列出文中提到的所有产品名，公司名、核心数据。
标注其【真实性】：（例如：真实存在 / 存在偏差 / 完全捏造！）。
【逻辑与专业性毒舌批注】（Critical Review）
摘录文中出现逻辑硬伤、外行话或生搬硬套的原句。
主编批注：用犀利的语言指出为什么这句话在业内人士看来是错的或可笑的。
【主编勒令修改建议】（Actionable Feedback）
给出 1-3 条极其具体的修改指令（例如："把第三段那个捏造的买量数据删掉，换成近期 XX 游戏的真实大盘数据"，"第五段关于留存的分析太水了，补上对次留和七留的具体漏斗推演"）。""",

    # 修改智能体提示词
    "reviser": """你是一位资深的内容编辑，负责根据审稿意见对文章进行修改润色。

你的任务：
1. 仔细阅读审稿人给出的修改意见
2. 对照原始素材，确保修改后的内容忠实于素材
3. 逐条落实修改建议，提升文章质量
4. 保持原文的优点和风格

重要：
- 不得删除素材中的重要信息
- 不得添加素材中没有的内容
- 修改要有针对性，不是大段重写"""
}

# prompts.json 文件路径
PROMPTS_FILE = "prompts.json"


def load_prompts() -> Dict[str, Any]:
    """从文件加载提示词配置，支持多编辑角色"""
    if os.path.exists(PROMPTS_FILE):
        try:
            with open(PROMPTS_FILE, 'r', encoding='utf-8') as f:
                loaded = json.load(f)
                # 合并：使用加载的配置，但保留默认的编辑器选项结构
                result = DEFAULT_PROMPTS.copy()
                if "editor_options" in loaded:
                    result["editor_options"] = loaded["editor_options"]
                if "default_editor" in loaded:
                    result["default_editor"] = loaded["default_editor"]
                if "reviewer" in loaded:
                    result["reviewer"] = loaded["reviewer"]
                if "reviser" in loaded:
                    result["reviser"] = loaded["reviser"]
                return result
        except Exception as e:
            print(f"加载 prompts.json 失败: {e}")
    return DEFAULT_PROMPTS.copy()


def save_prompts(prompts: Dict[str, Any]) -> bool:
    """保存提示词配置到文件"""
    try:
        with open(PROMPTS_FILE, 'w', encoding='utf-8') as f:
            json.dump(prompts, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        print(f"保存 prompts.json 失败: {e}")
        return False


# ============== 素材提取 ==============
def extract_web_content(url: str) -> Dict[str, Any]:
    """从网页URL提取内容"""
    try:
        import requests
        from bs4 import BeautifulSoup

        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }

        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()
        response.encoding = response.apparent_encoding or 'utf-8'

        soup = BeautifulSoup(response.text, 'html.parser')

        # 提取标题
        title = ""
        if soup.title:
            title = soup.title.string
        else:
            h1 = soup.find('h1')
            if h1:
                title = h1.get_text(strip=True)

        # 移除噪声标签
        for tag in soup(['script', 'style', 'nav', 'footer', 'header', 'aside']):
            tag.decompose()

        # 获取文本内容
        text = soup.get_text(separator='\n', strip=True)

        # 清理
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        lines = [line for line in lines if len(line) > 15]
        text = '\n'.join(lines)

        if not text or len(text) < 100:
            return {"success": False, "error": "提取的内容太少"}

        return {
            "success": True,
            "type": "web",
            "title": title.strip() if title else "未获取到标题",
            "content": text,
            "url": url
        }
    except Exception as e:
        return {"success": False, "error": f"网页提取失败: {str(e)}"}


def extract_youtube_content(url: str) -> Dict[str, Any]:
    """从YouTube视频提取字幕"""
    try:
        import requests

        # 提取视频ID
        video_id = None
        patterns = [
            r'(?:youtube\.com/watch\?v=|youtu\.be/|youtube\.com/embed/)([a-zA-Z0-9_-]{11})',
            r'youtube\.com/shorts/([a-zA-Z0-9_-]{11})'
        ]
        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                video_id = match.group(1)
                break

        if not video_id:
            return {"success": False, "error": "无法解析YouTube链接"}

        # 方法1：尝试使用 yt-dlp 或 youtube-transcript-api
        try:
            from youtube_transcript_api import YouTubeTranscriptApi
            transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=['zh-Hans', 'zh-Hant', 'en'])
            text = ' '.join([item['text'] for item in transcript])
            return {
                "success": True,
                "type": "youtube",
                "title": f"YouTube视频内容",
                "content": text,
                "url": url,
                "video_id": video_id
            }
        except:
            pass

        # 方法2：使用 Jina AI 提取
        try:
            jina_url = f"https://r.jina.ai/{url}"
            headers = {"Accept": "application/json"}
            response = requests.get(jina_url, headers=headers, timeout=30)
            if response.status_code == 200:
                data = response.json()
                return {
                    "success": True,
                    "type": "youtube",
                    "title": data.get("title", "YouTube视频内容"),
                    "content": data.get("content", ""),
                    "url": url,
                    "video_id": video_id
                }
        except:
            pass

        return {"success": False, "error": "无法获取字幕，请尝试其他链接"}

    except Exception as e:
        return {"success": False, "error": f"YouTube提取失败: {str(e)}"}


def extract_content(url: str) -> Dict[str, Any]:
    """统一的内容提取入口"""
    url_lower = url.lower()

    if 'youtube.com' in url_lower or 'youtu.be' in url_lower:
        return extract_youtube_content(url)
    else:
        return extract_web_content(url)


# ============== AI 调用 ==============
async def call_ai_api(
    messages: List[Dict],
    api_key: str,
    model: str = "gemini-2.5-flash"
) -> str:
    """调用 AI API"""
    url = f"{DEER_API_BASE_URL}/v1/chat/completions"

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": model,
        "messages": messages,
        "temperature": 0.7,
        "max_tokens": 8192
    }

    async with httpx.AsyncClient(timeout=180.0) as client:
        response = await client.post(url, headers=headers, json=payload)
        response.raise_for_status()
        data = response.json()
        return data["choices"][0]["message"]["content"]


def generate_draft(content: str, api_key: str, model: str, editor_prompt: str) -> Dict[str, Any]:
    """Step 2: 编辑生成初稿"""
    import asyncio

    user_prompt = f"""请根据以下原始素材，创作一篇微信公众号文章：

## 原始素材
{content}

## 要求
1. 标题要吸引人，包含关键词
2. 结构清晰，使用小标题分隔章节
3. 内容丰富有价值
4. 语言流畅，适合公众号阅读
5. 适当使用表情符号增加趣味性
6. 长度适中（800-2000字）
7. 结尾添加标签

请直接输出文章内容，不需要额外说明。"""

    try:
        result = asyncio.run(call_ai_api([
            {"role": "system", "content": editor_prompt},
            {"role": "user", "content": user_prompt}
        ], api_key, model))

        return {"success": True, "content": result}
    except Exception as e:
        return {"success": False, "error": str(e)}


def review_draft(
    original_content: str,
    draft: str,
    api_key: str,
    model: str,
    reviewer_prompt: str
) -> Dict[str, Any]:
    """Step 3: 审稿人审查（事实核查）"""
    import asyncio

    user_prompt = f"""## 原始素材（唯一的事实基准）
请严格对照此素材检查文章的准确性：
---
{original_content[:3000]}
---

## 待审稿的文章
---
{draft}
---

## 审稿任务
1. **事实核查**：文章内容是否忠实于原始素材？是否有虚构或夸大？
2. **完整性检查**：文章是否涵盖了素材的主要观点？
3. **可读性评估**：语言是否流畅？结构是否清晰？
4. **给出评分**：准确性、完整性、可读性各1-10分

## 输出格式（必须严格按JSON格式）
{{
    "accuracy_score": 1-10,
    "completeness_score": 1-10,
    "readability_score": 1-10,
    "overall_score": 平均分,
    "accuracy_issues": ["具体的事实错误列表"],
    "suggestions": ["修改建议列表"],
    "strengths": ["文章优点列表"],
    "can_publish": true/false
}}"""

    try:
        result = asyncio.run(call_ai_api([
            {"role": "system", "content": reviewer_prompt},
            {"role": "user", "content": user_prompt}
        ], api_key, model))

        # 解析 JSON
        try:
            review = json.loads(result)
            return {"success": True, "review": review}
        except:
            # 尝试从文本中提取 JSON
            import re
            json_match = re.search(r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}', result, re.DOTALL)
            if json_match:
                try:
                    review = json.loads(json_match.group(0))
                    return {"success": True, "review": review}
                except:
                    pass
            # 如果还是失败，返回原始文本作为错误信息
            return {"success": False, "error": f"审稿结果格式解析失败，返回内容: {result[:500]}..."}
    except Exception as e:
        return {"success": False, "error": str(e)}


def revise_draft(
    original_content: str,
    draft: str,
    review: Dict,
    api_key: str,
    model: str,
    reviser_prompt: str
) -> Dict[str, Any]:
    """Step 4: 根据审稿意见修改文章"""
    import asyncio

    suggestions_text = "\n".join([f"- {s}" for s in review.get("suggestions", [])])

    user_prompt = f"""## 原始素材（必须忠实于素材，不得添加虚构内容）
---
{original_content[:3000]}
---

## 原文章
---
{draft}
---

## 审稿修改意见
{suggestions_text}

## 修改任务
请根据上述修改意见，对文章进行修改润色。
- 逐条落实修改建议
- 不得添加素材中没有的内容
- 不得删除素材中的重要信息
- 保持原文的优点和风格

请直接输出修改后的完整文章。"""

    try:
        result = asyncio.run(call_ai_api([
            {"role": "system", "content": reviser_prompt},
            {"role": "user", "content": user_prompt}
        ], api_key, model))

        return {"success": True, "content": result}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ============== 导出功能 ==============
def export_to_docx(content: str, title: str) -> bytes:
    """导出为 Word 文档"""
    try:
        from docx import Document
        from docx.shared import Pt
        from io import BytesIO

        doc = Document()
        doc.add_heading(title, 0)

        # 添加内容（处理 Markdown 格式）
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            else:
                # 处理 Markdown 强调
                line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                line = re.sub(r'\*(.+?)\*', r'\1', line)
                doc.add_paragraph(line)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except ImportError:
        return None


def send_to_feishu(content: str, webhook_url: str) -> Dict[str, Any]:
    """发送内容到飞书群"""
    try:
        import requests

        payload = {
            "msg_type": "text",
            "content": {
                "text": f"📝 微信公众号文章内容：\n\n{content[:4000]}"
            }
        }

        response = requests.post(webhook_url, json=payload, timeout=10)
        result = response.json()

        if result.get("code") == 0:
            return {"success": True}
        else:
            return {"success": False, "error": result.get("msg", "发送失败")}
    except Exception as e:
        return {"success": False, "error": str(e)}


def parse_article(text: str) -> Dict[str, Any]:
    """解析文章内容"""
    result = {"title": "", "body": "", "tags": []}

    # 提取标题
    title_match = re.search(r'^#\s+(.+)$', text, re.MULTILINE)
    if title_match:
        result["title"] = title_match.group(1).strip()

    title_match = re.search(r'^标题[：:]\s*(.+?)(?:\n|$)', text, re.MULTILINE)
    if title_match and not result["title"]:
        result["title"] = title_match.group(1).strip()

    # 提取标签
    tags_match = re.search(r'标签[：:]\s*\[?(.*?)\]?(?:\n|$)', text, re.MULTILINE)
    if tags_match:
        tags_text = tags_match.group(1)
        result["tags"] = [t.strip() for t in re.split(r'[,，]', tags_text) if t.strip()]

    # 提取正文
    body = re.sub(r'^#\s+.+?(?:\n|$)', '', text, flags=re.MULTILINE)
    body = re.sub(r'^标题[：:].+?(?:\n|$)', '', body, flags=re.MULTILINE)
    body = re.sub(r'^正文[：:]\s*', '', body)
    body = re.sub(r'^标签[：:].+?(?:\n|$)', '', body, flags=re.MULTILINE)
    body = re.sub(r'\n{3,}', '\n\n', body)
    result["body"] = body.strip()

    return result


# ============== Streamlit UI ==============
def main():
    st.set_page_config(
        page_title="微信公众号文章生成工作流",
        page_icon="📝",
        layout="wide"
    )

    # 加载提示词配置
    prompts = load_prompts()

    # 标题
    st.title("📝 微信公众号文章生成工作流")
    st.markdown("**多智能体协作：编辑 → 审稿（事实核查）→ 修改 → 定稿**")

    # ============== 侧边栏 ==============
    with st.sidebar:
        st.header("⚙️ 设置")

        # API 密钥
        api_key = st.text_input(
            "DeerAPI 密钥",
            value=DEFAULT_API_KEY or "",
            type="password",
            help="从 https://api.deerapi.com 获取"
        )

        if not api_key:
            st.warning("⚠️ 请输入 DeerAPI 密钥")

        # 模型选择
        model_options = [m["name"] for m in AI_MODELS]
        selected_model_name = st.selectbox("AI 模型", model_options)
        selected_model = next(m["id"] for m in AI_MODELS if m["name"] == selected_model_name)

        st.divider()

        # 编辑角色选择
        st.header("🎯 编辑角色")

        # 获取编辑角色列表
        editor_options = prompts.get("editor_options", DEFAULT_PROMPTS.get("editor_options", []))
        if not editor_options:
            editor_options = DEFAULT_PROMPTS.get("editor_options", [])

        # 角色选择下拉框
        default_editor = prompts.get("default_editor", DEFAULT_PROMPTS.get("default_editor", "发行主编"))
        editor_names = [e["name"] for e in editor_options]
        selected_editor_name = st.selectbox(
            "选择编辑角色",
            options=editor_names,
            index=next((i for i, e in enumerate(editor_options) if e["id"] == default_editor), 0)
        )

        # 显示选中角色的描述
        selected_editor = next((e for e in editor_options if e["name"] == selected_editor_name), None)
        if selected_editor:
            st.caption(selected_editor.get("description", ""))

        # 保存选中的编辑器ID到session state
        st.session_state.selected_editor_id = selected_editor["id"] if selected_editor else default_editor

        # 显示选中角色的提示词
        st.session_state.editor_prompt = selected_editor.get("prompt", "") if selected_editor else ""

        # 查看详细提示词
        with st.expander("📋 查看角色提示词详情"):
            for editor in editor_options:
                st.markdown(f"**{editor['name']}**")
                st.caption(editor.get("description", ""))
                prompt_text = editor.get("prompt", "")[:300]
                if len(editor.get("prompt", "")) > 300:
                    prompt_text += "..."
                st.text(prompt_text)
                st.divider()

        st.divider()

        # 提示词管理
        st.header("🤖 提示词管理")

        # 加载/保存按钮
        col1, col2 = st.columns(2)
        with col1:
            if st.button("💾 保存", use_container_width=True):
                prompts_to_save = {
                    "editor_options": prompts.get("editor_options", DEFAULT_PROMPTS.get("editor_options", [])),
                    "default_editor": st.session_state.get("selected_editor_id", "发行主编"),
                    "reviewer": st.session_state.get("reviewer_prompt_input", prompts.get("reviewer", DEFAULT_PROMPTS["reviewer"])),
                    "reviser": st.session_state.get("reviser_prompt_input", prompts.get("reviser", DEFAULT_PROMPTS["reviser"]))
                }
                if save_prompts(prompts_to_save):
                    st.success("✅ 已保存到 prompts.json")
                else:
                    st.error("❌ 保存失败")

        with col2:
            if st.button("🔄 重置", use_container_width=True):
                prompts = DEFAULT_PROMPTS.copy()
                save_prompts(prompts)
                st.rerun()

        # 审稿人提示词编辑
        with st.expander("🔍 审稿人提示词（毒舌主编）", expanded=False):
            reviewer_prompt = st.text_area(
                "审稿 Prompt",
                value=prompts.get("reviewer", DEFAULT_PROMPTS["reviewer"]),
                height=200,
                key="reviewer_prompt_input"
            )

        # 修改智能体提示词
        with st.expander("✂️ 修改智能体提示词", expanded=False):
            reviser_prompt = st.text_area(
                "修改 Prompt",
                value=prompts.get("reviser", DEFAULT_PROMPTS["reviser"]),
                height=150,
                key="reviser_prompt_input"
            )

        st.divider()

        # 导出设置
        st.header("📤 导出设置")

        feishu_webhook = st.text_input(
            "飞书 Webhook URL",
            value=FEISHU_WEBHOOK_URL or "",
            placeholder="可选，发送文章到飞书群",
            help="在飞书群设置中添加自定义机器人获取"
        )

    # ============== 主内容区 ==============
    tab1, tab2, tab3, tab4 = st.tabs(["📥 输入", "🕷️ 素材", "⚙️ 工作流", "📄 结果"])

    # ---------- Tab 1: 输入 ----------
    with tab1:
        st.header("📥 输入内容")

        col1, col2 = st.columns([3, 1])

        with col1:
            source_url = st.text_input(
                "🔗 内容链接",
                placeholder="输入文章链接、YouTube视频链接...",
                help="支持网页文章、YouTube视频"
            )

            source_type = st.radio(
                "📌 内容类型",
                ["🌐 网页文章", "📺 YouTube视频", "📝 手动输入"],
                horizontal=True,
                help="选择内容来源类型"
            )

            content_type = "web" if "网页" in source_type else ("youtube" if "YouTube" in source_type else "manual")

        with col2:
            st.markdown("### 💡 提示")
            st.markdown("""
            - 支持常见新闻网站
            - 支持 YouTube 视频（自动提取字幕）
            - 也可以手动粘贴内容
            """)

        if content_type == "manual":
            manual_content = st.text_area(
                "📝 手动输入内容",
                height=200,
                placeholder="在此粘贴文章内容..."
            )
            if manual_content:
                st.session_state.manual_content = manual_content
        else:
            if st.button("🔍 提取内容", type="primary"):
                if not source_url:
                    st.error("请输入链接地址")
                elif not api_key:
                    st.error("请先输入 DeerAPI 密钥")
                else:
                    with st.spinner("正在提取内容..."):
                        result = extract_content(source_url)

                        if result["success"]:
                            st.session_state.extracted_content = result
                            st.success(f"✅ {result['type'].upper()} 内容提取成功！")
                        else:
                            st.error(f"❌ 提取失败: {result.get('error', '未知错误')}")

    # ---------- Tab 2: 素材预览 ----------
    with tab2:
        st.header("🕷️ 素材池")

        content_to_show = None

        if "extracted_content" in st.session_state:
            content_to_show = st.session_state.extracted_content
        elif "manual_content" in st.session_state:
            content_to_show = {
                "success": True,
                "type": "manual",
                "title": "手动输入内容",
                "content": st.session_state.manual_content,
                "url": ""
            }

        if content_to_show:
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(f"**📌 标题：** {content_to_show.get('title', '无')}")
            with col2:
                st.markdown(f"**📎 类型：** `{content_to_show.get('type', 'unknown')}`")

            if content_to_show.get('url'):
                st.markdown(f"**🔗 来源：** {content_to_show['url']}")

            st.divider()

            word_count = len(content_to_show.get("content", ""))
            st.markdown(f"**📊 字数：** 约 {word_count} 字")

            st.text_area(
                "素材内容",
                value=content_to_show.get("content", ""),
                height=400,
                disabled=True
            )
        else:
            st.info("👆 请先在「输入」标签页提取或输入内容")

    # ---------- Tab 3: 工作流 ----------
    with tab3:
        st.header("⚙️ 多智能体工作流")

        if not content_to_show:
            st.info("👆 请先在「输入」标签页提取内容")
        else:
            if not api_key:
                st.error("⚠️ 请先输入 DeerAPI 密钥")
            else:
                st.markdown("""
                ### 工作流程
                ```
                素材 → Step 2: 编辑生成初稿 → Step 3: 审稿人审查 → Step 4: 修改润色 → 定稿
                ```
                """)

                if st.button("▶️ 启动多智能体工作流", type="primary", use_container_width=True):
                    progress_bar = st.progress(0)
                    status_text = st.empty()

                    try:
                        original_content = content_to_show["content"]

                        # Step 2: 生成初稿
                        status_text.text("📝 Step 2/4：编辑智能体正在生成初稿...")
                        progress_bar.progress(20)
                        st.info("正在调用编辑智能体...")

                        draft_result = generate_draft(
                            original_content, api_key, selected_model,
                            st.session_state.get("editor_prompt", "")
                        )

                        if not draft_result["success"]:
                            st.error(f"❌ 初稿生成失败: {draft_result['error']}")
                            st.stop()

                        draft = draft_result["content"]
                        st.session_state.draft = draft
                        st.success("✅ 初稿生成完成")

                        with st.expander("📄 查看初稿", expanded=False):
                            st.markdown(draft)

                        # Step 3: 审稿审查
                        status_text.text("🔍 Step 3/4：审稿人正在审查（事实核查）...")
                        progress_bar.progress(50)
                        st.info("审稿人正在核查事实...")

                        review_result = review_draft(
                            original_content, draft, api_key, selected_model,
                            st.session_state.get("reviewer_prompt_input", "")
                        )

                        if not review_result["success"]:
                            st.error(f"❌ 审稿失败: {review_result['error']}")
                            st.stop()

                        review = review_result.get("review", {})
                        st.session_state.review = review
                        st.success(f"✅ 审稿完成，总分：{review.get('overall_score', '?')}/10")

                        # 显示审稿结果
                        col1, col2, col3, col4 = st.columns(4)
                        with col1:
                            st.metric("准确性", review.get("accuracy_score", "-"))
                        with col2:
                            st.metric("完整性", review.get("completeness_score", "-"))
                        with col3:
                            st.metric("可读性", review.get("readability_score", "-"))
                        with col4:
                            score = review.get("overall_score", 0)
                            st.metric("总分", score,
                                     delta_color="normal" if score >= 8 else "inverse")

                        if review.get("suggestions"):
                            st.markdown("**💡 修改建议：**")
                            for s in review["suggestions"]:
                                st.markdown(f"- {s}")

                        if review.get("strengths"):
                            st.markdown("**✨ 优点：**")
                            for s in review["strengths"]:
                                st.markdown(f"- {s}")

                        # Step 4: 修改润色
                        if review.get("overall_score", 0) < 8:
                            status_text.text("✂️ Step 4/4：根据意见修改文章...")
                            progress_bar.progress(75)
                            st.info("正在根据审稿意见修改文章...")

                            revise_result = revise_draft(
                                original_content, draft, review, api_key, selected_model,
                                st.session_state.get("reviser_prompt_input", "")
                            )

                            if revise_result["success"]:
                                final_article = revise_result["content"]
                                st.session_state.final_article = final_article
                                st.success("✅ 修改完成！")
                            else:
                                st.warning(f"⚠️ 修改失败，使用初稿: {revise_result.get('error', '')}")
                                st.session_state.final_article = draft
                        else:
                            st.success("🎉 文章质量优秀，无需修改！")
                            st.session_state.final_article = draft

                        # 完成
                        status_text.text("✅ 工作流完成！")
                        progress_bar.progress(100)
                        st.balloons()

                    except Exception as e:
                        st.error(f"❌ 执行出错: {str(e)}")

    # ---------- Tab 4: 结果 ----------
    with tab4:
        st.header("📄 生成结果")

        if "final_article" not in st.session_state:
            st.info("👆 请先执行工作流生成文章")
        else:
            article_data = parse_article(st.session_state.final_article)

            # 文章标题
            st.subheader(f"📌 {article_data['title'] or '生成的标题'}")

            # 标签
            if article_data["tags"]:
                st.markdown("**标签：** " + " ".join([f"`{t}`" for t in article_data["tags"]]))

            # 显示文章内容
            st.markdown(article_data["body"])

            st.divider()

            # 导出选项
            col1, col2, col3 = st.columns(3)

            with col1:
                # Markdown 下载
                md_content = f"""# {article_data['title']}

{article_data['body']}

---
标签：{', '.join(article_data['tags']) if article_data['tags'] else '无'}
生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
"""
                st.download_button(
                    "📥 下载 Markdown",
                    md_content,
                    file_name=f"{article_data['title'] or 'article'}.md",
                    mime="text/markdown",
                    use_container_width=True
                )

            with col2:
                # Word 文档下载
                try:
                    from docx import Document
                    from io import BytesIO

                    doc = Document()
                    doc.add_heading(article_data['title'] or '文章标题', 0)

                    for para in article_data['body'].split('\n'):
                        para = para.strip()
                        if para:
                            if para.startswith('## '):
                                doc.add_heading(para[3:], level=2)
                            elif para.startswith('# '):
                                doc.add_heading(para[2:], level=1)
                            else:
                                doc.add_paragraph(para)

                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)

                    st.download_button(
                        "📄 下载 Word",
                        buffer.getvalue(),
                        file_name=f"{article_data['title'] or 'article'}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                except ImportError:
                    st.caption("需要 python-docx 库")

            with col3:
                # 复制到剪贴板 - 使用文本框方式
                st.text_area("📋 复制全文（选中后 Ctrl+C）", value=md_content, height=100, key="copy_text")

            # 发送飞书
            if feishu_webhook:
                st.divider()
                if st.button("🚀 发送到飞书群", use_container_width=True):
                    result = send_to_feishu(article_data["body"], feishu_webhook)
                    if result["success"]:
                        st.success("✅ 已发送到飞书群！")
                    else:
                        st.error(f"❌ 发送失败: {result.get('error', '')}")

            # 审稿结果
            if "review" in st.session_state:
                st.divider()
                st.subheader("🔍 审稿结果")

                review = st.session_state["review"]
                if isinstance(review, dict):
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("准确性", review.get("accuracy_score", "-"))
                    with col2:
                        st.metric("完整性", review.get("completeness_score", "-"))
                    with col3:
                        st.metric("可读性", review.get("readability_score", "-"))
                    with col4:
                        score = review.get("overall_score", 0)
                        st.metric("总分", score,
                                 delta_color="normal" if score >= 8 else "inverse")


if __name__ == "__main__":
    main()
